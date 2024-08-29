from docx import Document
import os

def save_docx_by_pages(input_docx_path, output_folder):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)

    doc = Document(input_docx_path)
    full_text = []
    current_text = []
    doc_index = 1

    for para in doc.paragraphs:
        current_text.append(para.text)
        if '' in para.text:  # Form feed character  usually denotes a new page
            full_text.append(current_text)
            current_text = []
            doc_index += 1

    # In case the document doesn't end with a form feed
    if current_text:
        full_text.append(current_text)

    for i, page_content in enumerate(full_text):
        new_doc = Document()
        for para_text in page_content:
            new_doc.add_paragraph(para_text)
        new_doc.save(os.path.join(output_folder, f"page_{i+1}.docx"))

if __name__ == "__main__":
    import sys
    if len(sys.argv) != 3:
        print("Usage: python split_docx.py <input_docx_path> <output_folder>")
    else:
        input_docx_path = sys.argv[1]
        output_folder = sys.argv[2]
        save_docx_by_pages(input_docx_path, output_folder)
        print(f"Pages have been split and saved in {output_folder}")
